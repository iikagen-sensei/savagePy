"""
docx_generator.py
Renderiza un template .docx con docxtpl + prepara el contexto de datos.
Los templates usan sintaxis Jinja2 estándar: {{ }}, {% %}, {% endfor %}.
"""
import io
from pathlib import Path


def render_docx_template(template_path: Path, context: dict) -> bytes:
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def build_powers_context(data: list, titulo: str = "Poderes", view_name: str = "") -> dict:
    # Datos en crudo — usa {{ p.name }}, {{ p.cost }}, {{ p.description }}...
    # Datos agrupados y formateados — usa {{ entrada.nombre }}, {{ entrada.coste_label }}...
    grupos = {}
    for p in data:
        cat = p.get("rank_name") or "Sin rango"
        grupos.setdefault(cat, []).append(p)

    bloques = []
    for cat, items in grupos.items():
        entradas = []
        for p in items:
            nombre = p.get("name", "")
            if p.get("name_original"):
                nombre += f" ({p['name_original']})"
            ref = p.get("reference_book") or {}
            fuente = ref.get("title", "")
            if p.get("page_no"):
                fuente += f" p.{p['page_no']}"
            mods = []
            for m in (p.get("modifiers") or []):
                t = m.get("title", "")
                if m.get("cost"):
                    t += f" ({m['cost']})"
                mods.append({"titulo": t, "descripcion": m.get("description", "")})
            entradas.append({
                # Campos formateados (labels)
                "nombre":          nombre,
                "coste_label":     f"Coste: {p['cost']} PP" if p.get("cost") else "",
                "alcance_label":   f"Alcance: {p.get('range_roh') or p.get('range', '')}" if (p.get("range_roh") or p.get("range")) else "",
                "duracion_label":  f"Duración: {p['duration']}" if p.get("duration") else "",
                "fuente":          fuente,
                "descripcion":     p.get("description", ""),
                "modificadores":   mods,
                # Campos en crudo (nombres tal cual vienen de NocoDB)
                **p,
            })
        bloques.append({"categoria": cat, "entradas": entradas})

    return {"titulo": titulo, "view_name": view_name, "bloques": bloques, "powers": data}


def build_edges_context(data: list, titulo: str = "Ventajas", view_name: str = "") -> dict:
    # Datos en crudo — usa {{ e.title }}, {{ e.requirements }}, {{ e.description }}...
    # Datos agrupados y formateados — usa {{ entrada.nombre }}, {{ entrada.requisitos_label }}...
    grupos = {}
    for e in data:
        cat = e.get("type") or e.get("rank_name") or "General"
        grupos.setdefault(cat, []).append(e)

    bloques = []
    for cat, items in grupos.items():
        entradas = []
        for e in items:
            nombre = e.get("title") or e.get("name") or ""
            if e.get("name_original"):
                nombre += f" ({e['name_original']})"
            ref = e.get("reference_book") or {}
            fuente = ref.get("title", "")
            if e.get("page_no"):
                fuente += f" p.{e['page_no']}"
            entradas.append({
                # Campos formateados (labels)
                "nombre":           nombre,
                "requisitos_label": f"Requisitos: {e['requirements']}" if e.get("requirements") else "",
                "fuente":           fuente,
                "descripcion":      e.get("description", ""),
                # Campos en crudo
                **e,
            })
        bloques.append({"categoria": cat, "entradas": entradas})

    return {"titulo": titulo, "view_name": view_name, "bloques": bloques, "edges": data}


def build_hindrances_context(data: list, titulo: str = "Desventajas", view_name: str = "") -> dict:
    # Datos en crudo — usa {{ h.Nombre }}, {{ h.type }}, {{ h.description }}...
    # Datos agrupados y formateados — usa {{ entrada.nombre }}, {{ entrada.fuente }}...
    grupos = {}
    for h in data:
        cat = h.get("type") or "General"
        grupos.setdefault(cat, []).append(h)

    bloques = []
    for cat, items in grupos.items():
        entradas = []
        for h in items:
            nombre = h.get("Nombre") or h.get("name") or h.get("title") or ""
            if h.get("name_original"):
                nombre += f" ({h['name_original']})"
            ref = h.get("reference_book") or {}
            fuente = ref.get("title", "")
            if h.get("page_no"):
                fuente += f" p.{h['page_no']}"
            entradas.append({
                # Campos formateados (labels)
                "nombre":      nombre,
                "fuente":      fuente,
                "descripcion": h.get("description", ""),
                # Campos en crudo
                **h,
            })
        bloques.append({"categoria": cat, "entradas": entradas})

    return {"titulo": titulo, "view_name": view_name, "bloques": bloques, "hindrances": data}


CONTEXT_BUILDERS = {
    "powers":     build_powers_context,
    "edges":      build_edges_context,
    "hindrances": build_hindrances_context,
}


def build_rules_docx(data: list, titulo: str = "Reglas", view_name: str = "") -> bytes:
    """
    Genera un .docx completo con todas las reglas convirtiendo el contenido
    Markdown a Word con pypandoc. No usa plantilla — genera el documento por código
    para preservar el formato enriquecido (negritas, tablas, listas, títulos).
    """
    import pypandoc
    import tempfile
    import os
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Título del documento
    title = doc.add_heading(titulo, level=0)
    if view_name:
        sub = doc.add_paragraph(view_name)
        sub.runs[0].italic = True

    for rule in data:
        # Encabezado de la regla
        name = rule.get("name", "")
        name_original = rule.get("name_original")
        heading_text = f"{name} ({name_original})" if name_original else name
        doc.add_heading(heading_text, level=1)

        # Metadatos en una línea
        meta_parts = []
        source = rule.get("source")
        if source:
            meta_parts.append(source)
        ref = rule.get("reference_book") or {}
        ref_title = ref.get("title") if isinstance(ref, dict) else None
        if ref_title:
            page = rule.get("page_no")
            meta_parts.append(f"{ref_title} p.{page}" if page else ref_title)
        if meta_parts:
            meta = doc.add_paragraph(" · ".join(meta_parts))
            meta.runs[0].italic = True

        # Descripción breve
        description = rule.get("description")
        if description:
            desc = doc.add_paragraph(description)
            desc.runs[0].bold = True

        # Contenido Markdown → docx via pypandoc
        content_md = rule.get("content")
        if content_md:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                pypandoc.convert_text(
                    content_md, "docx", format="md",
                    outputfile=tmp_path,
                    extra_args=["--standalone"]
                )
                # Extraer párrafos del docx generado e insertarlos
                tmp_doc = Document(tmp_path)
                for elem in tmp_doc.element.body:
                    doc.element.body.append(elem)
            finally:
                os.unlink(tmp_path)

        # Separador entre reglas
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

